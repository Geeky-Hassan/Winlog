from typing import Annotated,Optional
from datetime import datetime
from fastapi import FastAPI,Depends, HTTPException,status,File,UploadFile,Form,Body
from fastapi.staticfiles import StaticFiles
from fastapi.security import OAuth2PasswordBearer
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import and_, func, desc
from sqlalchemy.orm import Session
from database import session_local,Users,Brags
from globe.back_func import *
import bcrypt
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from fastapi.responses import StreamingResponse
import requests
import logging
import google.generativeai as genai
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.graphics.shapes import Drawing, Line, String, Rect
from reportlab.lib.utils import ImageReader
from reportlab.lib.enums import TA_CENTER
from PIL import Image as PILImage
import io

load_dotenv()  # This loads the variables from .env

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")
app:FastAPI = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=['http://localhost:3000'],
    allow_credentials=True,
    allow_methods=['*'],
    allow_headers=['*'])
app.mount('/uploads/uploaded_brags', StaticFiles(directory="uploaded_brags"), name="uploads")

def get_db():
    db:Session = session_local()
    try:
        yield db
    finally:
        db.close()

@app.get('/')
def index():
    return {'message': 'Welcome to the FastAPI REST API!'}

@app.post('/register')
def register(data: dict, db: Session = Depends(get_db)):
    try:
        # Check if the email already exists
        existing_user = db.query(Users).filter(Users.user_mail == data['email']).first()
        if existing_user:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Email already registered")

        print(f"data: {data['password']}")
        hashed_pass = bcrypt.hashpw(data['password'].encode(), bcrypt.gensalt()) # encoding password here
        
        decoded = hashed_pass.decode('utf-8')   # decoding password here
        
        user = Users(user_mail=data['email'], user_name=data['username'], user_pass=decoded) 
        db.add(user)
        db.commit()
        db.refresh(user)
        token = generate_token(data)
        return {"message": "User created successfully", "status": status.HTTP_201_CREATED, "authToken": token}
    except Exception as e:
        print(e)
        return HTTPException(500, "User Already Exist")


@app.post('/login')
def login(data: dict, db: Session = Depends(get_db)):
    try:
        # Get the user by email
        user = db.query(Users).filter(Users.user_mail == data['email']).first()
        
        # If user not found, raise an exception
        if user is None:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid email or password")
        
        # Check the password using bcrypt
        if not bcrypt.checkpw(data['password'].encode(), user.user_pass.encode()):
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid email or password")
        
        user_info = {
            "id": user.user_id,
            "email": user.user_mail,
            "is_admin": user.is_admin,
            "disabled": user.disabled
        }
        # Generate the token
        token = generate_token(user_info)
        
        # Return the token
        return {"message": "Login successful", "authToken": token, "status": status.HTTP_200_OK}
    
    except HTTPException as he:
        # Re-raise HTTP exceptions
        raise he
    except Exception as e:
        logger.error(f"Error during login: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Error logging in")


async def get_current_user(token: Annotated[str, Depends(oauth2_scheme)]):
    token =decode_token(token)
    return token


@app.get("/g_brags")
async def get_brag(
    token: Annotated[dict, Depends(get_current_user)],
    include_soft_deleted: bool = False,
    db: Session = Depends(get_db)
):
    email = token['email']
    user_id = db.query(Users).filter(Users.user_mail == email).first().user_id 
    query = db.query(Brags).filter(Brags.users == user_id)
    
    if not include_soft_deleted:
        query = query.filter(Brags.is_soft_deleted == False)
    
    # Sort by updated_time first, then by created_time, both in descending order
    brags = query.order_by(desc(Brags.updated_time), desc(Brags.created_time)).all()
    
    # Strip whitespace from brag names before returning
    for brag in brags:
        brag.brag_name = brag.brag_name.strip()
    
    return {"brags": brags}

@app.post("/add_brag")
async def add_brag(
    token: Annotated[dict, Depends(get_current_user)],
    title: str = Form(...),
    desc: str = Form(...),
    tags: str = Form(...),  # Change this from list to str
    designation: Optional[str] = Form(None),
    start_date: str = Form(...),
    end_date: Optional[str] = Form(None),
    img: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
): 
    email= token['email']
    
    try:
        user = db.query(Users).filter(Users.user_mail == email).first()
    
        
        if user is None:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized user")
        
        img_path = None
        if img is not None:
            img_path = generate_image_path(img.filename, img)
        
        # Get current datetime
        current_time = datetime.now()

        brag = Brags(
            brag_name=title.strip(),  # Strip whitespace here
            brag_desc=desc,
            brag_designation=designation,
            brag_tags=tags,  # This should now be a string
            brag_img=img_path,
            brag_start_date=start_date,
            brag_end_date=end_date,
            created_time=current_time,
            updated_time=current_time,
            users=user.user_id
        )
        
        # Try to add, commit, and refresh the brag object
        try:
            db.add(brag)
            db.commit()
            db.refresh(brag)
        except Exception as e:
            # Rollback the session in case of an error
            db.rollback()
            print(f"Error while adding brag: {e}")  # Log the error
            raise HTTPException(detail="Could not create brag at this moment", status_code=status.HTTP_400_BAD_REQUEST)
        
        return {"detail": "Brag created successfully!", "status_code": status.HTTP_201_CREATED, "brag_id":brag.brag_id}
    
    except Exception as e:
        print(e)  # Log the error for debugging
        raise HTTPException(detail="Could not create brag at this moment", status_code=status.HTTP_400_BAD_REQUEST)

@app.put("/u_brag")
async def update_brag(
    token: Annotated[dict, Depends(get_current_user)],
    pTitle: str = Form(...),
    title: Optional[str] = Form(None),
    desc: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    designation: Optional[str] = Form(None),
    start_date: Optional[str] = Form(None),
    end_date: Optional[str] = Form(None),
    img: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
): 
    email = token['email']

    try:
        user = db.query(Users).filter(Users.user_mail == email).first()
        
        if user is None:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized user")
        
        old_brag = db.query(Brags).filter(and_(
            Brags.users == user.user_id,
            func.lower(func.trim(Brags.brag_name)) == func.lower(pTitle.strip())
        )).first()

        if old_brag is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Old brag not found")

        # Create a new brag object with updated fields
        new_brag = Brags(
            brag_name=title if title is not None else old_brag.brag_name,
            brag_desc=desc if desc is not None else old_brag.brag_desc,
            brag_tags=tags if tags is not None else old_brag.brag_tags,
            brag_designation=designation if designation is not None else old_brag.brag_designation,
            brag_start_date=start_date if start_date is not None else old_brag.brag_start_date,
            brag_end_date=end_date if end_date is not None else old_brag.brag_end_date,
            brag_img=old_brag.brag_img,  # Keep the old image by default
            created_time=old_brag.created_time,
            updated_time=datetime.now(),
            users=user.user_id
        )

        # Handle image update
        if img:
            new_img_path = generate_image_path(img.filename, img)
            new_brag.brag_img = new_img_path

        # Add the new brag to the database
        db.add(new_brag)
        db.flush()  # This will assign an ID to new_brag

        # Update the old brag
        old_brag.is_soft_deleted = True
        old_brag.brag_next = new_brag.brag_id
        new_brag.brag_prev = old_brag.brag_id

        db.commit()

        return {"detail": "Brag updated successfully!", "status_code": status.HTTP_200_OK}
    
    except Exception as e:
        db.rollback()
        print(e)  # Log the error for debugging
        raise HTTPException(detail="Could not update brag at this moment", status_code=status.HTTP_400_BAD_REQUEST)

@app.delete("/d_brags")
async def delete_brag(
    token: Annotated[dict, Depends(get_current_user)],
    brag_title: str = Body(..., embed=True),
    db: Session = Depends(get_db)
):
    email = token['email']
    
    try:
        user = db.query(Users).filter(Users.user_mail == email).first()
        
        if user is None:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized user")
        
        brag = db.query(Brags).filter(and_(
            Brags.users == user.user_id,
            func.lower(func.trim(Brags.brag_name)) == func.lower(brag_title.strip())
        )).first()
        
        if brag is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Brag with title '{brag_title}' not found for this user")
        
        # Delete associated image if it exists
        if brag.brag_img:
            img_path = os.path.join("uploaded_brags", brag.brag_img)
            if os.path.exists(img_path):
                os.remove(img_path)
        
        # Delete all versions of this brag
        db.query(Brags).filter(Brags.brag_id == brag.brag_id).delete()
        db.query(Brags).filter(Brags.brag_prev == brag.brag_id).delete()
        db.query(Brags).filter(Brags.brag_next == brag.brag_id).delete()
        
        db.commit()
        
        return {"detail": "Brag and all its versions deleted successfully!", "status_code": status.HTTP_200_OK}
    
    except HTTPException as he:
        db.rollback()
        raise he
    except Exception as e:
        db.rollback()
        print(f"Error deleting brag: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Could not delete brag: {str(e)}")

# ----------------------------
# Revert Brag to previous state
# ----------------------------
@app.post("/r_brags")
async def revert_brag(
    token: Annotated[dict, Depends(get_current_user)],
    brag_title: str = Body(..., embed=True),
    db: Session = Depends(get_db)
):
    email = token['email']
    print(f"Attempting to revert brag for user: {email}, brag title: {brag_title}")
    
    try:
        user = db.query(Users).filter(Users.user_mail == email).first()
        if user is None:
            print(f"User not found: {email}")
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized user")
        
        # Find the current brag by title, strip whitespace
        current_brag = db.query(Brags).filter(and_(
            Brags.users == user.user_id,
            func.lower(func.trim(Brags.brag_name)) == func.lower(brag_title.strip())
        )).first()
        
        if current_brag is None:
            print(f"Brag not found: {brag_title}")
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Brag not found")
        
        print(f"Current brag found: {current_brag.brag_id}")
        
        # Check if the current brag has a previous version
        if current_brag.brag_prev is None:
            print(f"No previous version for brag: {current_brag.brag_id}")
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No previous version to revert to")

        # Find the previous brag using the prev field
        previous_brag = db.query(Brags).filter(Brags.brag_id == current_brag.brag_prev).first()

        if previous_brag is None:
            print(f"Previous brag not found: {current_brag.brag_prev}")
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Previous version not found")

        print(f"Previous brag found: {previous_brag.brag_id}")

        # Revert the previous brag by setting is_soft_deleted to False
        db.query(Brags).filter(Brags.brag_id == previous_brag.brag_id).update({
            Brags.is_soft_deleted: False,
            Brags.brag_next: None 
        })

        # Delete the current brag
        db.delete(current_brag)

        # Commit the changes
        db.commit()
        print("Revert successful")

        return {"detail": "Brag reverted to previous version successfully!", "status_code": status.HTTP_200_OK}
    
    except HTTPException as he:
        print(f"HTTPException: {he.detail}")
        raise he
    except Exception as e:
        print(f"Error reverting brag: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Could not revert brag: {str(e)}")

@app.post('/google-login')
async def google_login(token: dict, db: Session = Depends(get_db)):
    try:
        print("Received token:", token)  # Log the received token
        
        # Create credentials using the access token
        credentials = Credentials(token['token'])
        
        # Build the Google People API service
        service = build('people', 'v1', credentials=credentials)
        
        # Call the People API
        person = service.people().get(resourceName='people/me', personFields='names,emailAddresses').execute()
        
        # Extract user info
        email = person['emailAddresses'][0]['value']
        name = person['names'][0]['displayName']
        print(f"User info - Email: {email}, Name: {name}")  # Log user info

        # Check if user exists in your database
        user = db.query(Users).filter(Users.user_mail == email).first()
        print("Existing user:", user)  # Log if user exists

        if not user:
            # Create a new user if they don't exist
            user = Users(user_mail=email, user_name=name, user_pass="google_auth")
            db.add(user)
            db.commit()
            db.refresh(user)
            print("New user created:", user)  # Log new user creation

        # Generate JWT token
        token = generate_token({"email": email, "id": user.user_id})
        print("Generated JWT token:", token)  # Log generated token

        return {"message": "Google login successful", "authToken": token, "status": status.HTTP_200_OK}

    except Exception as e:
        print("Exception:", str(e))  # Log general exception
        raise HTTPException(status_code=500, detail=f"Could not process Google login: {str(e)}")

@app.get("/download_brags_word")
async def download_brags_word(token: Annotated[dict, Depends(get_current_user)], db: Session = Depends(get_db)):
    email = token['email']
    user_id = db.query(Users).filter(Users.user_mail == email).first().user_id 
    
    # Fetch only the latest version of each brag, ordered by updated_time
    latest_brags = db.query(Brags).filter(
        Brags.users == user_id,
        Brags.brag_next == None,  # This ensures we get only the latest version
        Brags.is_soft_deleted == False
    ).order_by(desc(Brags.updated_time)).all()

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add logo
    logo_path = os.path.join("uploaded_brags", "logo.png")
    if os.path.exists(logo_path):
        document.add_picture(logo_path, width=Inches(3))

    document.add_heading('Your Brag Timeline', 0)

    def add_brag_card(brag, document):
        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        
        # Add image
        if brag.brag_img:
            img_path = os.path.abspath(os.path.join(brag.brag_img))
            if os.path.exists(img_path):
                try:
                    cell.add_paragraph().add_run().add_picture(img_path, width=Inches(3))
                except Exception as e:
                    cell.add_paragraph(f"Error loading image: {str(e)}")
            else:
                cell.add_paragraph("Image not found")

        # Add brag details
        title = cell.add_paragraph(brag.brag_name)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Heading 2'
        
        start_date = brag.brag_start_date or "N/A"
        end_date = brag.brag_end_date or "Present"
        date = cell.add_paragraph(f"Date: {start_date} - {end_date}")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        desc = brag.brag_desc[:100] + "..." if brag.brag_desc and len(brag.brag_desc) > 100 else (brag.brag_desc or "No description")
        description = cell.add_paragraph(desc)
        description.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        tags = brag.brag_tags.split(',') if brag.brag_tags else ["No tags"]
        tag_para = cell.add_paragraph()
        tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tag in tags:
            tag_run = tag_para.add_run(f" {tag.strip()} ")
            tag_run.font.highlight_color = WD_COLOR_INDEX.PINK
        
        # Set border color to primary
        set_cell_border(cell, 'FF8C0054', sz=12)  # Primary color

    def set_cell_border(cell, color, sz):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for border in ['top', 'left', 'bottom', 'right']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(sz))
            element.set(qn('w:color'), color)
            tcBorders.append(element)

    for brag in latest_brags:
        add_brag_card(brag, document)
        
        # Add timeline connector
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('|')
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x8C, 0x00, 0x54)  # Primary color

    # Add footer
    section = document.sections[-1]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "© 2023 Winlog. All rights reserved."
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx_file = BytesIO()
    document.save(docx_file)
    docx_file.seek(0)

    return StreamingResponse(
        iter([docx_file.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=brags_timeline.docx"}
    )

# Configure the Gemini AI
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

@app.get("/suggestions")
async def get_suggestions(query: str):
    try:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(f"Generate 5 creative and concise title suggestions based on this query: {query}. Return only the titles, one per line, without numbering.")
        suggestions = [suggestion.strip() for suggestion in response.text.split('\n') if suggestion.strip()]
        return {"suggestions": suggestions[:5]}  # Ensure we only return up to 5 suggestions
    except Exception as e:
        logger.error(f"Error generating suggestions: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating suggestions: {str(e)}")

@app.get("/hashtag-suggestions")
async def get_hashtag_suggestions(title: str):
    try:
        model = genai.GenerativeModel('gemini-pro')
        logger.info(f"Generating hashtags for title: {title}")
        response = model.generate_content(f"Generate 5 relevant and trendy hashtags for this title: {title}. Return only the hashtags without the '#' symbol, one per line, without numbering.")
        logger.info(f"Raw response from Gemini: {response.text}")
        hashtags = [hashtag.strip() for hashtag in response.text.split('\n') if hashtag.strip()]
        logger.info(f"Processed hashtags: {hashtags}")
        return {"hashtags": hashtags[:5]}  # Ensure we only return up to 5 hashtags
    except Exception as e:
        logger.error(f"Error generating hashtag suggestions: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating hashtag suggestions: {str(e)}")

@app.get("/download_brags_pdf")
async def download_brags_pdf(token: Annotated[dict, Depends(get_current_user)], db: Session = Depends(get_db)):
    email = token['email']
    user_id = db.query(Users).filter(Users.user_mail == email).first().user_id 
    
    # Fetch only the latest version of each brag, ordered by updated_time
    latest_brags = db.query(Brags).filter(
        Brags.users == user_id,
        Brags.brag_next == None,  # This ensures we get only the latest version
        Brags.is_soft_deleted == False
    ).order_by(desc(Brags.updated_time)).all()

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles using Tailwind root colors
    primary_color = HexColor("#8C0054")
    secondary_color = HexColor("#10B981")
    text_color = HexColor("#000000")
    bg_color = HexColor("#F3F4F6")

    styles['Heading1'].fontSize = 24
    styles['Heading1'].alignment = TA_CENTER
    styles['Heading1'].spaceAfter = 12
    styles['Heading1'].textColor = primary_color
    styles.add(ParagraphStyle(name='BragTitle', fontSize=16, textColor=primary_color, spaceAfter=6))
    styles.add(ParagraphStyle(name='BragInfo', fontSize=10, textColor=text_color, spaceAfter=6))
    styles.add(ParagraphStyle(name='BragDesc', fontSize=12, spaceAfter=12, textColor=text_color))
    styles.add(ParagraphStyle(name='Footer', fontSize=10, alignment=TA_CENTER, textColor=HexColor("#FFFFFF")))
    styles.add(ParagraphStyle(name='TagStyle', fontSize=8, textColor=HexColor("#FFFFFF"), backColor=primary_color))

    # Add logo
    logo_path = os.path.join("uploaded_brags", "logo.png")
    if os.path.exists(logo_path):
        img = Image(logo_path, width=3*inch, height=2*inch)
        story.append(img)

    story.append(Paragraph("Your Brag Timeline", styles['Heading1']))
    story.append(Spacer(1, 24))

    def create_brag_card(brag):
        card_content = []
        
        # Add image (if exists) with max height of 200px
        if brag.brag_img:
            img_path = os.path.abspath(os.path.join(brag.brag_img))
            logger.info(f"Attempting to load image from: {img_path}")
            if os.path.exists(img_path):
                try:
                    with open(img_path, 'rb') as img_file:
                        img = Image(img_file)
                        aspect = img.imageWidth / float(img.imageHeight)
                        img.drawHeight = 1.5*inch  # ~150px
                        img.drawWidth = 1.5*inch * aspect
                        card_content.append(img)
                    logger.info(f"Successfully loaded image: {img_path}")
                except Exception as e:
                    logger.error(f"Error loading image {img_path}: {str(e)}")
                    card_content.append(Paragraph(f"Error loading image", styles['BragInfo']))
            else:
                logger.warning(f"Image file not found: {img_path}")
                card_content.append(Paragraph(f"Image not found", styles['BragInfo']))
        
        card_content.append(Paragraph(brag.brag_name, styles['BragTitle']))
        card_content.append(Paragraph(f"Date: {brag.brag_start_date} - {brag.brag_end_date}", styles['BragInfo']))
        
        desc = brag.brag_desc[:100] + "..." if brag.brag_desc and len(brag.brag_desc) > 100 else (brag.brag_desc or "No description")
        card_content.append(Paragraph(desc, styles['BragDesc']))
        
        tags = brag.brag_tags.split(',') if brag.brag_tags else ["No tags"]
        tag_paragraphs = [Paragraph(tag.strip(), styles['TagStyle']) for tag in tags]
        card_content.append(Table([tag_paragraphs], colWidths=[1*inch]*len(tag_paragraphs), style=[
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        return Table([[content] for content in card_content], colWidths=[6*inch], style=[
            ('BACKGROUND', (0, 0), (-1, -1), bg_color),
            ('TEXTCOLOR', (0, 0), (-1, -1), text_color),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOX', (0, 0), (-1, -1), 2, primary_color),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ])

    for brag in latest_brags:
        story.append(create_brag_card(brag))
        story.append(Spacer(1, 0.2*inch))

    # Footer
    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(primary_color)
        canvas.rect(0, 0, doc.pagesize[0], 0.5*inch, fill=True)
        footer_text = Paragraph("© 2023 Winlog. All rights reserved.", styles['Footer'])
        w, h = footer_text.wrap(doc.pagesize[0] - 0.5*inch, doc.bottomMargin)
        footer_text.drawOn(canvas, 0.25*inch, 0.25*inch)
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    buffer.seek(0)

    return StreamingResponse(
        iter([buffer.getvalue()]),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=brags_timeline.pdf"}
    )

if __name__ == '__main__':
    import uvicorn
    uvicorn.run("server:app", reload=True)
